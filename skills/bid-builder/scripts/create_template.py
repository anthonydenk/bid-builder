#!/usr/bin/env python3
"""Create the branded, client-customizable Bid Builder DOCX template."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def main() -> None:
    output = Path(__file__).resolve().parent.parent / "assets/proposal-template.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.5)
    header = section.header.paragraphs[0]
    header.text = "PARTITIONS & ACCESSORIES CO.  /  COMMERCIAL PROPOSAL"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor(216, 138, 61)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("1220 South Pasadena, Mesa, AZ 85210  •  480-969-6606  •  ROC # CR60 110352")
    doc.add_paragraph("This body is replaced by the proposal generator.")
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
