#!/usr/bin/env python3
"""Deterministic CAD-export inspection and proposal generation for Bid Builder."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import zipfile
from collections import OrderedDict
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET


MONEY = Decimal("0.01")
TOLERANCE = Decimal("0.02")
REQUIRED_HEADERS = {
    "type",
    "item",
    "quantity",
    "description",
    "net cost",
    "unit profit amount",
    "unit price",
    "extended cost",
    "extended profit",
    "net price",
}
INTERNAL_TYPES = {"freight", "labor", "overhead", "travel", "textura", "resource"}
NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_PKG_REL = "http://schemas.openxmlformats.org/package/2006/relationships"


class BidBuilderError(RuntimeError):
    """Actionable input or environment error."""


def money(value: Decimal) -> Decimal:
    return value.quantize(MONEY, rounding=ROUND_HALF_UP)


def quantity_text(value: Decimal) -> str:
    if value == value.to_integral_value():
        return format(value, ".0f")
    return format(value.normalize(), "f")


def decimal(value: Any) -> Decimal:
    if value in (None, ""):
        return Decimal("0")
    try:
        return Decimal(str(value).replace(",", "").replace("$", "").strip())
    except (InvalidOperation, AttributeError) as exc:
        raise BidBuilderError(f"Expected a number, found {value!r}") from exc


def column_number(cell_ref: str) -> int:
    letters = re.match(r"[A-Z]+", cell_ref.upper())
    if not letters:
        return 0
    result = 0
    for char in letters.group(0):
        result = result * 26 + ord(char) - 64
    return result - 1


def normalized_header(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).lower()


def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    except KeyError:
        return []
    values: list[str] = []
    for item in root.findall(f"{{{NS_MAIN}}}si"):
        values.append("".join(node.text or "" for node in item.iter(f"{{{NS_MAIN}}}t")))
    return values


def _sheet_targets(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    workbook = ET.fromstring(archive.read("xl/workbook.xml"))
    rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    targets = {
        rel.attrib["Id"]: rel.attrib["Target"]
        for rel in rels.findall(f"{{{NS_PKG_REL}}}Relationship")
    }
    result: list[tuple[str, str]] = []
    sheets = workbook.find(f"{{{NS_MAIN}}}sheets")
    if sheets is None:
        return []
    for sheet in sheets:
        rel_id = sheet.attrib.get(f"{{{NS_REL}}}id", "")
        target = targets.get(rel_id, "")
        if target.startswith("/"):
            path = target.lstrip("/")
        elif target.startswith("xl/"):
            path = target
        else:
            path = f"xl/{target.lstrip('/')}"
        result.append((sheet.attrib.get("name", "Sheet"), path))
    return result


def _cell_value(cell: ET.Element, shared: list[str]) -> Any:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.iter(f"{{{NS_MAIN}}}t"))
    value_node = cell.find(f"{{{NS_MAIN}}}v")
    if value_node is None or value_node.text is None:
        return ""
    value = value_node.text
    if cell_type == "s":
        try:
            return shared[int(value)]
        except (ValueError, IndexError):
            return value
    if cell_type in {"str", "b"}:
        return value
    try:
        return Decimal(value)
    except InvalidOperation:
        return value


def _sheet_rows(archive: zipfile.ZipFile, target: str, shared: list[str]) -> list[tuple[int, list[Any]]]:
    root = ET.fromstring(archive.read(target))
    rows: list[tuple[int, list[Any]]] = []
    for row in root.findall(f".//{{{NS_MAIN}}}sheetData/{{{NS_MAIN}}}row"):
        values: list[Any] = []
        for cell in row.findall(f"{{{NS_MAIN}}}c"):
            idx = column_number(cell.attrib.get("r", "A1"))
            while len(values) <= idx:
                values.append("")
            values[idx] = _cell_value(cell, shared)
        rows.append((int(row.attrib.get("r", len(rows) + 1)), values))
    return rows


@dataclass
class ExportRow:
    source: str
    sheet: str
    row_number: int
    type: str
    item: str
    quantity: Decimal
    description: str
    net_cost: Decimal
    unit_profit: Decimal
    unit_price: Decimal
    extended_cost: Decimal
    extended_profit: Decimal
    net_price: Decimal
    category: str

    def to_dict(self) -> dict[str, Any]:
        result = self.__dict__.copy()
        for key, value in result.items():
            if isinstance(value, Decimal):
                result[key] = str(money(value))
        return result


def read_export(path: Path) -> tuple[list[ExportRow], list[str]]:
    if path.suffix.lower() == ".xls":
        raise BidBuilderError(f"{path.name}: legacy .xls is not supported; re-export it as .xlsx")
    if path.suffix.lower() != ".xlsx":
        raise BidBuilderError(f"{path.name}: expected an .xlsx workbook")
    if not path.exists():
        raise BidBuilderError(f"Workbook does not exist: {path}")

    parsed: list[ExportRow] = []
    warnings: list[str] = []
    with zipfile.ZipFile(path) as archive:
        shared = _read_shared_strings(archive)
        for sheet_name, target in _sheet_targets(archive):
            rows = _sheet_rows(archive, target, shared)
            header_index = None
            header_map: dict[str, int] = {}
            for index, (_, values) in enumerate(rows):
                candidate = {normalized_header(v): i for i, v in enumerate(values) if str(v).strip()}
                if REQUIRED_HEADERS.issubset(candidate):
                    header_index = index
                    header_map = candidate
                    break
            if header_index is None:
                warnings.append(f"{path.name}/{sheet_name}: no recognized export header; sheet skipped")
                continue

            def get(values: list[Any], name: str) -> Any:
                idx = header_map.get(name)
                return values[idx] if idx is not None and idx < len(values) else ""

            for row_number, values in rows[header_index + 1 :]:
                description = str(get(values, "description") or "").strip()
                item = str(get(values, "item") or "").strip()
                quantity = decimal(get(values, "quantity"))
                net_price = decimal(get(values, "net price"))
                if not description and not item and quantity == 0 and net_price == 0:
                    continue
                parsed.append(
                    ExportRow(
                        source=path.name,
                        sheet=sheet_name,
                        row_number=row_number,
                        type=str(get(values, "type") or "").strip(),
                        item=item,
                        quantity=quantity,
                        description=description,
                        net_cost=decimal(get(values, "net cost")),
                        unit_profit=decimal(get(values, "unit profit amount")),
                        unit_price=decimal(get(values, "unit price")),
                        extended_cost=decimal(get(values, "extended cost")),
                        extended_profit=decimal(get(values, "extended profit")),
                        net_price=net_price,
                        category=str(get(values, "category") or "").strip(),
                    )
                )
    if not parsed:
        raise BidBuilderError(f"{path.name}: no export rows were found")
    return parsed, warnings


def validate_row(row: ExportRow) -> list[str]:
    checks = [
        (row.quantity * row.net_cost, row.extended_cost, "quantity × net cost", "extended cost"),
        (row.quantity * row.unit_profit, row.extended_profit, "quantity × unit profit", "extended profit"),
        (row.net_cost + row.unit_profit, row.unit_price, "net cost + unit profit", "unit price"),
        (row.extended_cost + row.extended_profit, row.net_price, "extended cost + extended profit", "net price"),
    ]
    issues = []
    for calculated, stated, formula, field in checks:
        if abs(money(calculated) - money(stated)) > TOLERANCE:
            issues.append(
                f"{row.source}/{row.sheet} row {row.row_number}: {formula} = ${money(calculated):,.2f}, "
                f"but {field} is ${money(stated):,.2f}"
            )
    return issues


class Catalog:
    def __init__(self, path: Path):
        payload = json.loads(path.read_text(encoding="utf-8"))
        self.rules = payload.get("rules", [])

    def resolve(self, row: ExportRow) -> dict[str, Any]:
        public = row.type.strip().lower() not in INTERNAL_TYPES
        result = {"public": public, "manufacturer": None, "model": None, "rule_id": None}
        exact_rules = []
        regex_rules = []
        for rule in self.rules:
            match = rule.get("match", {})
            if match.get("item"):
                exact_rules.append(rule)
            else:
                regex_rules.append(rule)
        for rule in [*exact_rules, *regex_rules]:
            match = rule.get("match", {})
            item_match = not match.get("item") or row.item.casefold() == str(match["item"]).casefold()
            regex_match = not match.get("description_regex") or re.search(
                str(match["description_regex"]), row.description
            )
            if item_match and regex_match:
                result.update(
                    public=bool(rule.get("public", public)),
                    manufacturer=rule.get("manufacturer"),
                    model=rule.get("model") or row.item,
                    rule_id=rule.get("id"),
                )
                return result
        return result


def resolve_path(raw: str, job_path: Path) -> Path:
    path = Path(raw).expanduser()
    return path if path.is_absolute() else (job_path.parent / path).resolve()


def required_job_blockers(job: dict[str, Any]) -> list[str]:
    blockers = []
    for field in ("naming_line", "project", "proposal_number", "customer", "location"):
        if not str(job.get(field, "")).strip():
            blockers.append(f"Missing required job field: {field}")
    if job.get("grouping_confirmed") is not True:
        blockers.append("Spreadsheet grouping has not been explicitly confirmed")
    for field in ("proposal_date", "plan_date", "prepared_by"):
        if not str(job.get(field) or "").strip():
            blockers.append(f"Missing proposal fact: {field}")
    if not job.get("clauses_approved", False):
        blockers.append("Selected proposal clauses have not been approved by the company")
    if not job.get("sections"):
        blockers.append("No proposal sections were supplied")
    return blockers


def build_model(job: dict[str, Any], job_path: Path, catalog_path: Path, clauses_path: Path) -> dict[str, Any]:
    catalog = Catalog(catalog_path)
    clauses_payload = json.loads(clauses_path.read_text(encoding="utf-8"))
    profile_name = job.get("clauses_profile", "general_draft")
    profile = clauses_payload.get("profiles", {}).get(profile_name)
    blockers = required_job_blockers(job)
    if profile is None:
        blockers.append(f"Unknown clause profile: {profile_name}")
        profile = []

    sections = []
    grand_total = Decimal("0")
    seen_source_paths: set[str] = set()
    for section_index, section in enumerate(job.get("sections", []), start=1):
        title = str(section.get("title") or f"Section {section_index}")
        kind = str(section.get("kind") or "other").lower()
        if kind == "partitions":
            if not str(section.get("manufacturer") or "").strip():
                blockers.append(f"{title}: partition manufacturer is required")
            if not str(section.get("scope_summary") or "").strip():
                blockers.append(f"{title}: partition scope summary is required")
            if not str(section.get("installation_basis") or "").strip():
                blockers.append(f"{title}: partition installation basis is required")

        all_rows: list[ExportRow] = []
        section_warnings: list[str] = []
        source_files = []
        if not section.get("source_files"):
            blockers.append(f"{title}: no source files were supplied")
        for raw_path in section.get("source_files", []):
            path = resolve_path(str(raw_path), job_path)
            normalized_source = str(path)
            if normalized_source in seen_source_paths:
                blockers.append(f"Workbook was assigned more than once: {path.name}")
            seen_source_paths.add(normalized_source)
            rows, warnings = read_export(path)
            all_rows.extend(rows)
            section_warnings.extend(warnings)
            source_files.append(str(path))

        arithmetic_issues = [issue for row in all_rows for issue in validate_row(row)]
        blockers.extend(arithmetic_issues)
        source_total = money(sum((row.net_price for row in all_rows), Decimal("0")))
        grand_total += source_total

        products: OrderedDict[tuple[str, str, str], dict[str, Any]] = OrderedDict()
        visible_value = Decimal("0")
        unresolved: list[str] = []
        if kind != "partitions":
            for row in all_rows:
                resolved = catalog.resolve(row)
                if not resolved["public"]:
                    continue
                manufacturer = resolved.get("manufacturer") or "UNRESOLVED"
                model = resolved.get("model") or row.item
                if manufacturer == "UNRESOLVED":
                    unresolved.append(
                        f"{title}: manufacturer unresolved for {row.item or '(no item)'} — {row.description}"
                    )
                key = (manufacturer, str(model), row.description)
                if key not in products:
                    products[key] = {
                        "manufacturer": manufacturer,
                        "model": model,
                        "description": row.description,
                        "quantity": Decimal("0"),
                        "source_value": Decimal("0"),
                    }
                products[key]["quantity"] += row.quantity
                products[key]["source_value"] += row.net_price
                visible_value += row.net_price
        blockers.extend(dict.fromkeys(unresolved))
        public_products = []
        for product in products.values():
            public_products.append(
                {
                    **product,
                    "quantity": quantity_text(product["quantity"]),
                    "source_value": str(money(product["source_value"])),
                }
            )

        hidden_allowance = money(source_total - visible_value)
        sections.append(
            {
                "title": title,
                "group_title": section.get("group_title") or title,
                "kind": kind,
                "manufacturer": section.get("manufacturer"),
                "scope_summary": section.get("scope_summary"),
                "installation_basis": section.get("installation_basis") or "Furnished & Installed",
                "source_files": source_files,
                "source_row_count": len(all_rows),
                "source_total": str(source_total),
                "public_source_value": str(money(visible_value)),
                "hidden_allowance": str(hidden_allowance),
                "products": public_products,
                "warnings": section_warnings,
                "arithmetic_issues": arithmetic_issues,
            }
        )

    unique_blockers = list(dict.fromkeys(blockers))
    return {
        "schema_version": 1,
        "generated_on": date.today().isoformat(),
        "client_ready": not unique_blockers,
        "draft_reason": None if not unique_blockers else "Resolve all blockers before issue.",
        "job": {
            key: job.get(key)
            for key in (
                "naming_line",
                "project",
                "proposal_number",
                "customer",
                "location",
                "proposal_date",
                "plan_date",
                "prepared_by",
                "attention",
                "email",
                "phone",
            )
        },
        "sections": sections,
        "grand_total": str(money(grand_total)),
        "clause_profile": profile_name,
        "clauses": profile,
        "blockers": unique_blockers,
    }


def write_reconciliation(model: dict[str, Any], path: Path) -> None:
    lines = [
        f"# Reconciliation — {model['job'].get('proposal_number') or 'unassigned'}",
        "",
        f"**Status:** {'CLIENT READY' if model['client_ready'] else 'DRAFT — NOT FOR ISSUE'}",
        "",
        f"**Naming line:** {model['job'].get('naming_line') or 'MISSING'}",
        "",
    ]
    for section in model["sections"]:
        lines.extend(
            [
                f"## {section['title']}",
                "",
                f"- Source files: {', '.join(Path(p).name for p in section['source_files'])}",
                f"- Parsed rows: {section['source_row_count']}",
                f"- Public product value: ${Decimal(section['public_source_value']):,.2f}",
                f"- Hidden allowance: ${Decimal(section['hidden_allowance']):,.2f}",
                f"- Source section total: ${Decimal(section['source_total']):,.2f}",
                "",
            ]
        )
        if section["arithmetic_issues"]:
            lines.append("Arithmetic issues:")
            lines.extend(f"- {issue}" for issue in section["arithmetic_issues"])
            lines.append("")
        if section["warnings"]:
            lines.append("Warnings:")
            lines.extend(f"- {warning}" for warning in section["warnings"])
            lines.append("")
    lines.extend(["## Grand total", "", f"**${Decimal(model['grand_total']):,.2f}**", ""])
    lines.append("## Blockers")
    lines.append("")
    lines.extend(f"- {item}" for item in model["blockers"]) if model["blockers"] else lines.append("- None")
    lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell: Any, color: str = "000000", size: int = 4) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell: Any, top: int = 35, start: int = 70, bottom: int = 35, end: int = 70) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table: Any, widths_inches: list[float]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    table.autofit = False
    total_twips = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def _repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_run_font(run: Any, size: float, bold: bool = False, italic: bool = False, color: str = "000000") -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _paragraph_spacing(paragraph: Any, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _display_text(value: Any) -> str:
    return str(value or "").replace("\u2013", "-").replace("\u2014", "-")


def _long_date(value: Any) -> str:
    text = str(value or "").strip()
    try:
        return datetime.fromisoformat(text).strftime("%A, %B %d, %Y")
    except ValueError:
        return text or "Pending"


def _short_date(value: Any) -> str:
    text = str(value or "").strip()
    try:
        parsed = datetime.fromisoformat(text)
        return f"{parsed.month}/{parsed.day}/{str(parsed.year)[2:]}"
    except ValueError:
        return text or "Pending"


def write_docx(model: dict[str, Any], path: Path, template_path: Path | None = None) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise BidBuilderError("DOCX generation requires python-docx: pip install -r requirements.txt") from exc

    doc = Document(str(template_path)) if template_path and template_path.exists() else Document()
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.34)
    section.right_margin = Inches(0.34)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.16)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    styles["Normal"].paragraph_format.line_spacing = 1.0

    doc.core_properties.title = f"Proposal {model['job'].get('proposal_number') or 'Unassigned'}"
    doc.core_properties.subject = _display_text(model["job"].get("project"))
    doc.core_properties.author = "Partitions & Accessories Co."

    header = section.header.paragraphs[0]
    header.text = ""

    footer = section.footer.paragraphs[0]
    footer.text = ""
    clauses = model.get("clauses") if isinstance(model.get("clauses"), dict) else {}
    footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    footer_run = footer.add_run(_display_text(clauses.get("footer_terms", "")))
    _set_run_font(footer_run, 5.5)
    _paragraph_spacing(footer, line=0.92)

    profile_path = Path(__file__).resolve().parent.parent / "references/company-profile.json"
    company = json.loads(profile_path.read_text(encoding="utf-8"))
    content_width = 7.82

    masthead = doc.add_table(rows=1, cols=2)
    masthead.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(masthead, [4.35, 3.47])
    logo_path = Path(__file__).resolve().parent.parent / "assets/company-logo.jpg"
    logo_paragraph = masthead.cell(0, 0).paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path.exists():
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(3.35))
    else:
        _set_run_font(logo_paragraph.add_run(company.get("legal_name", "Partitions & Accessories Co.")), 16, True, color="0B4B96")
    contact_cell = masthead.cell(0, 1)
    contact_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(contact_cell, start=110)
    contact_lines = [
        company.get("address", ""),
        f"{company.get('phone', '')}    {company.get('website', '')}",
        f"{company.get('license', '')}    Tax ID#{company.get('tax_id', '')}",
    ]
    contact_paragraph = contact_cell.paragraphs[0]
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, line in enumerate(contact_lines):
        run = contact_paragraph.add_run(("\n" if index else "") + line)
        _set_run_font(run, 9)
    _paragraph_spacing(contact_paragraph, line=1.06)
    _set_cell_border(contact_cell, color="FFFFFF", size=0)
    left_border = contact_cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    if left_border is not None:
        left = left_border.find(qn("w:left"))
        if left is not None:
            left.set(qn("w:color"), "5BC0DE")
            left.set(qn("w:sz"), "5")

    validity = doc.add_paragraph()
    _paragraph_spacing(validity, before=1, after=5)
    _set_run_font(validity.add_run(f"Quote is good for {clauses.get('quote_valid_days', 30)} Days from date of proposal"), 9.5, True)
    if not model["client_ready"]:
        validity.alignment = WD_ALIGN_PARAGRAPH.LEFT
        draft_run = validity.add_run("    DRAFT - NOT FOR ISSUE")
        _set_run_font(draft_run, 9, True, color="C00000")

    metadata = doc.add_table(rows=1, cols=3)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(metadata, [2.45, 2.7, 2.67])
    date_p = metadata.cell(0, 0).paragraphs[0]
    _set_run_font(date_p.add_run(_long_date(model["job"].get("proposal_date"))), 11)
    proposal_p = metadata.cell(0, 1).paragraphs[0]
    proposal_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(proposal_p.add_run(f"Proposal # {model['job'].get('proposal_number') or 'UNASSIGNED'}"), 12.5, True)
    plans_cell = metadata.cell(0, 2)
    _set_cell_border(plans_cell, color="1F4E79", size=6)
    _set_cell_margins(plans_cell, top=55, start=95, bottom=55, end=95)
    plans_p = plans_cell.paragraphs[0]
    _set_run_font(plans_p.add_run("Per Plans Dated:\n"), 8.5, True, color="FF0000")
    _set_run_font(plans_p.add_run(_short_date(model["job"].get("plan_date"))), 9, True, color="FF0000")

    prepared = doc.add_paragraph()
    _paragraph_spacing(prepared, before=5, after=6)
    _set_run_font(prepared.add_run("Prepared by:    "), 9.5, True)
    _set_run_font(prepared.add_run(_display_text(model["job"].get("prepared_by") or "Pending")), 9.5, True)

    customer = doc.add_table(rows=3, cols=4)
    customer.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(customer, [0.78, 2.85, 1.27, 2.92])
    customer_values = [
        ("TO:", model["job"].get("customer"), "JOB:", model["job"].get("project")),
        ("ATTN:", model["job"].get("attention"), "LOCATION:", str(model["job"].get("location") or "").upper()),
        ("EMAIL:", model["job"].get("email"), "PHONE:", model["job"].get("phone")),
    ]
    for row, values in zip(customer.rows, customer_values):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.text = ""
            _set_cell_border(cell)
            _set_cell_margins(cell, top=20, bottom=20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(_display_text(value))
            _set_run_font(run, 9.2, bold=index in (0, 2))

    payment = doc.add_table(rows=1, cols=1)
    payment.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(payment, [content_width])
    pay_cell = payment.cell(0, 0)
    _set_cell_border(pay_cell, color="FF0000", size=14)
    _set_cell_margins(pay_cell, top=45, start=80, bottom=40, end=80)
    pay_p = pay_cell.paragraphs[0]
    _set_run_font(pay_p.add_run(clauses.get("payment_heading", "NON CREDIT CUSTOMERS:") + " "), 9.3, True)
    _set_run_font(pay_p.add_run(clauses.get("payment_intro", "")), 9.3)
    _paragraph_spacing(pay_p)
    for index, term in enumerate(clauses.get("payment_terms", []), start=1):
        item = pay_cell.add_paragraph()
        item.paragraph_format.left_indent = Inches(0.28)
        item.paragraph_format.first_line_indent = Inches(-0.18)
        _paragraph_spacing(item)
        _set_run_font(item.add_run(f"{index}.  {_display_text(term)}"), 8.7)

    last_group = None
    for section_model in model["sections"]:
        group_title = _display_text(section_model.get("group_title") or section_model["title"])
        if group_title != last_group:
            group = doc.add_paragraph()
            _paragraph_spacing(group, before=3, after=1)
            group_run = group.add_run(group_title)
            _set_run_font(group_run, 11.5, True)
            group_run.font.highlight_color = WD_COLOR_INDEX.RED
            last_group = group_title

        if section_model["kind"] == "partitions":
            heading = doc.add_paragraph()
            _paragraph_spacing(heading, after=0)
            _set_run_font(
                heading.add_run(f"Toilet Partitions By: {_display_text(section_model.get('manufacturer') or 'Pending')}"),
                10.5,
                True,
            )
            scope = doc.add_paragraph()
            _paragraph_spacing(scope, after=0)
            _set_run_font(scope.add_run(_display_text(section_model.get("scope_summary") or "Pending")), 10, True, True)
        else:
            heading = doc.add_paragraph()
            _paragraph_spacing(heading, after=1)
            _set_run_font(heading.add_run("Toilet Accessories:"), 10.5, True)

        if section_model["products"]:
            items = doc.add_table(rows=1, cols=4)
            items.alignment = WD_TABLE_ALIGNMENT.LEFT
            _set_table_widths(items, [1.48, 1.34, 2.8, 2.2])
            _repeat_table_header(items.rows[0])
            for cell, label in zip(items.rows[0].cells, ["Qty.", "No.", "Description", "Manufacturer Name"]):
                cell.text = ""
                _set_cell_border(cell)
                _set_cell_margins(cell, top=26, start=65, bottom=26, end=65)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_run_font(cell.paragraphs[0].add_run(label), 8.5, True)
            for product in section_model["products"]:
                cells = items.add_row().cells
                values = [
                    product["quantity"],
                    product["model"],
                    product["description"],
                    product["manufacturer"],
                ]
                for index, (cell, value) in enumerate(zip(cells, values)):
                    cell.text = ""
                    _set_cell_border(cell)
                    _set_cell_margins(cell, top=22, start=65, bottom=22, end=65)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    if index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _set_run_font(paragraph.add_run(_display_text(value)), 8.1)

        amount = doc.add_paragraph()
        _paragraph_spacing(amount, before=1, after=2)
        basis = _display_text(section_model.get("installation_basis") or "Furnished & Installed")
        amount_run = amount.add_run(f"{basis}: ${Decimal(section_model['source_total']):,.2f}")
        _set_run_font(amount_run, 10.5, True)
        amount_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_page_break()
    tax = doc.add_paragraph()
    tax.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(tax, before=2, after=9)
    tax_run = tax.add_run(_display_text(clauses.get("tax_note", "*Pricing excludes taxes*")))
    _set_run_font(tax_run, 10.5, True)
    tax_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    escalation = doc.add_paragraph()
    _paragraph_spacing(escalation, after=5)
    escalation_run = escalation.add_run(_display_text(clauses.get("escalation_heading", "Proposal Expiration & Potential Escalation")))
    _set_run_font(escalation_run, 10.5, True)
    escalation_run.underline = True
    for paragraph_text in clauses.get("escalation_paragraphs", []):
        paragraph = doc.add_paragraph()
        _paragraph_spacing(paragraph, after=5, line=1.08)
        _set_run_font(paragraph.add_run(_display_text(paragraph_text)), 9.2)

    exclusion_intro = doc.add_paragraph()
    exclusion_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(exclusion_intro, before=3, after=5)
    exclusion_run = exclusion_intro.add_run(_display_text(clauses.get("exclusion_intro", "")))
    _set_run_font(exclusion_run, 9.6, True, color="FF0000")
    exclusion_run.underline = True
    exclusion_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for exclusion in clauses.get("exclusions", []):
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.42)
        bullet.paragraph_format.first_line_indent = Inches(-0.2)
        _paragraph_spacing(bullet, after=1)
        _set_run_font(bullet.add_run(_display_text(exclusion)), 8.7, True, color="FF0000")

    warning = doc.add_table(rows=1, cols=1)
    warning.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(warning, [7.15])
    warning_cell = warning.cell(0, 0)
    _set_cell_shading(warning_cell, "FF0000")
    _set_cell_margins(warning_cell, top=38, bottom=38)
    warning_p = warning_cell.paragraphs[0]
    warning_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(warning_p.add_run(_display_text(clauses.get("price_warning", ""))), 9.2, True, color="FFFFFF")

    note = doc.add_paragraph()
    _paragraph_spacing(note, before=6, after=3, line=1.0)
    _set_run_font(note.add_run("NOTE: "), 7.6, True)
    _set_run_font(note.add_run(_display_text(clauses.get("partition_note", ""))), 7.6)
    _set_run_font(note.add_run(" Exclusions: "), 7.6, True)
    _set_run_font(note.add_run(_display_text(clauses.get("additional_exclusions", ""))), 7.6)

    order = doc.add_paragraph()
    _paragraph_spacing(order, before=4, after=9)
    _set_run_font(order.add_run(_display_text(clauses.get("order_instruction", ""))), 9.2, True)

    signature = doc.add_table(rows=1, cols=3)
    signature.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(signature, [3.45, 3.0, 1.37])
    signature_values = [
        "PRINTED NAME/TITLE: __________________________",
        "Signature: __________________________",
        "Date: __________",
    ]
    for cell, value in zip(signature.rows[0].cells, signature_values):
        cell.text = ""
        _set_cell_margins(cell, start=0, end=35)
        _set_run_font(cell.paragraphs[0].add_run(value), 8.6)

    doc.save(path)


def convert_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    completed = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        text=True,
        capture_output=True,
        check=False,
    )
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not pdf_path.exists():
        raise BidBuilderError(f"LibreOffice PDF conversion failed: {completed.stderr or completed.stdout}")
    return pdf_path


def command_inspect(paths: Iterable[str]) -> int:
    for raw in paths:
        path = Path(raw).expanduser().resolve()
        rows, warnings = read_export(path)
        issues = [issue for row in rows for issue in validate_row(row)]
        total = money(sum((row.net_price for row in rows), Decimal("0")))
        print(f"{path.name}: {len(rows)} rows, ${total:,.2f}")
        for message in [*warnings, *issues]:
            print(f"  - {message}")
    return 0


def command_build(job_arg: str, output_arg: str, no_pdf: bool) -> int:
    job_path = Path(job_arg).expanduser().resolve()
    output_dir = Path(output_arg).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    job = json.loads(job_path.read_text(encoding="utf-8"))
    skill_root = Path(__file__).resolve().parent.parent
    model = build_model(
        job,
        job_path,
        skill_root / "references/item-catalog.json",
        skill_root / "references/clauses.json",
    )
    model_path = output_dir / "proposal-model.json"
    model_path.write_text(json.dumps(model, indent=2) + "\n", encoding="utf-8")
    write_reconciliation(model, output_dir / "reconciliation.md")
    docx_path = output_dir / "proposal.docx"
    write_docx(model, docx_path, skill_root / "assets/proposal-template.docx")
    if not no_pdf:
        convert_pdf(docx_path, output_dir)
    print(f"Built {output_dir}")
    print(f"Status: {'CLIENT READY' if model['client_ready'] else 'DRAFT — NOT FOR ISSUE'}")
    print(f"Total: ${Decimal(model['grand_total']):,.2f}")
    if model["blockers"]:
        print(f"Blockers: {len(model['blockers'])}")
    return 0 if model["client_ready"] else 2


def parser() -> argparse.ArgumentParser:
    root = argparse.ArgumentParser(prog="bid-builder")
    subparsers = root.add_subparsers(dest="command", required=True)
    inspect_parser = subparsers.add_parser("inspect", help="Inspect and total CAD-exported XLSX files")
    inspect_parser.add_argument("workbooks", nargs="+")
    build_parser = subparsers.add_parser("build", help="Build a proposal package from a job JSON file")
    build_parser.add_argument("--job", required=True)
    build_parser.add_argument("--output-dir", required=True)
    build_parser.add_argument("--no-pdf", action="store_true")
    return root


def main(argv: list[str] | None = None) -> int:
    args = parser().parse_args(argv)
    try:
        if args.command == "inspect":
            return command_inspect(args.workbooks)
        return command_build(args.job, args.output_dir, args.no_pdf)
    except (BidBuilderError, json.JSONDecodeError, OSError, zipfile.BadZipFile) as exc:
        print(f"Bid Builder error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
