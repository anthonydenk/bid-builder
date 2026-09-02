import importlib.util
import json
import sys
import tempfile
import unittest
from decimal import Decimal
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "skills/bid-builder/scripts/bid_builder.py"
SPEC = importlib.util.spec_from_file_location("bid_builder", MODULE_PATH)
bid_builder = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
sys.modules["bid_builder"] = bid_builder
SPEC.loader.exec_module(bid_builder)


class BidBuilderTests(unittest.TestCase):
    def setUp(self):
        self.fixture_dir = ROOT / "tests/fixtures"
        self.catalog = ROOT / "skills/bid-builder/references/item-catalog.json"
        self.clauses = ROOT / "skills/bid-builder/references/clauses.json"

    def test_parser_preserves_source_total(self):
        rows, warnings = bid_builder.read_export(self.fixture_dir / "floor-1-accessories.xlsx")
        self.assertEqual([], warnings)
        self.assertEqual(2, len(rows))
        self.assertEqual(Decimal("340.00"), bid_builder.money(sum((r.net_price for r in rows), Decimal("0"))))
        self.assertEqual([], [issue for row in rows for issue in bid_builder.validate_row(row)])

    def test_skill_collects_complete_intake_packet_before_processing(self):
        skill = (ROOT / "skills/bid-builder/SKILL.md").read_text()
        for label in (
            "JOB / GROUPING:",
            "PROPOSAL DATE:",
            "PLANS DATED:",
            "PREPARED BY:",
            "PARTITIONS:",
            "PROPOSAL TERMS:",
        ):
            self.assertIn(label, skill)
        self.assertIn("before analyzing workbook contents", skill)

    def test_catalog_can_expose_product_encoded_as_freight(self):
        rows, _ = bid_builder.read_export(self.fixture_dir / "floor-2-accessories.xlsx")
        catalog = bid_builder.Catalog(self.catalog)
        kohler = next(row for row in rows if "KOHLER" in row.description)
        resolved = catalog.resolve(kohler)
        self.assertTrue(resolved["public"])
        self.assertEqual("Kohler", resolved["manufacturer"])
        self.assertEqual("14380-CP", resolved["model"])

    def test_build_groups_floors_and_reconciles_hidden_allowance(self):
        job = {
            "naming_line": "Example Medical, proposal Q10001, Example GC, Phoenix. Floors 1 and 2 go together.",
            "project": "Example Medical",
            "proposal_number": "Q10001",
            "customer": "Example GC",
            "location": "Phoenix",
            "grouping_confirmed": True,
            "proposal_date": "2026-08-28",
            "plan_date": "2026-08-14",
            "prepared_by": "Estimator Name",
            "clauses_profile": "general_draft",
            "clauses_approved": True,
            "sections": [
                {
                    "title": "Floors 1 and 2 — Accessories",
                    "kind": "accessories",
                    "source_files": [
                        str(self.fixture_dir / "floor-1-accessories.xlsx"),
                        str(self.fixture_dir / "floor-2-accessories.xlsx"),
                    ],
                },
                {
                    "title": "Floor 2 — Partitions",
                    "kind": "partitions",
                    "manufacturer": "Accurate",
                    "scope_summary": "(4) Stalls / Solid Phenolic / Overhead Braced",
                    "installation_basis": "Furnished & Installed",
                    "source_files": [str(self.fixture_dir / "floor-2-partitions.xlsx")],
                },
            ],
        }
        model = bid_builder.build_model(job, ROOT / "examples/job.example.json", self.catalog, self.clauses)
        self.assertTrue(model["client_ready"])
        self.assertEqual("4440.00", model["grand_total"])
        accessories = model["sections"][0]
        self.assertEqual("940.00", accessories["source_total"])
        self.assertEqual("160.00", accessories["hidden_allowance"])
        self.assertEqual({"ASI", "Bobrick", "Kohler"}, {p["manufacturer"] for p in accessories["products"]})

    def test_missing_intake_and_partition_scope_are_blockers(self):
        job = {
            "naming_line": "",
            "project": "Example",
            "proposal_number": "Q1",
            "customer": "GC",
            "location": "Mesa",
            "grouping_confirmed": False,
            "sections": [
                {
                    "title": "Partitions",
                    "kind": "partitions",
                    "source_files": [str(self.fixture_dir / "floor-2-partitions.xlsx")],
                }
            ],
        }
        model = bid_builder.build_model(job, ROOT / "job.json", self.catalog, self.clauses)
        self.assertFalse(model["client_ready"])
        self.assertTrue(any("naming_line" in item for item in model["blockers"]))
        self.assertTrue(any("partition manufacturer" in item for item in model["blockers"]))
        self.assertTrue(any("partition scope summary" in item for item in model["blockers"]))
        self.assertTrue(any("installation basis" in item for item in model["blockers"]))

    def test_outputs_model_reconciliation_and_docx(self):
        job_path = ROOT / "examples/job.example.json"
        job = json.loads(job_path.read_text())
        job["clauses_approved"] = True
        with tempfile.TemporaryDirectory() as temp:
            output = Path(temp)
            model = bid_builder.build_model(job, job_path, self.catalog, self.clauses)
            (output / "proposal-model.json").write_text(json.dumps(model))
            bid_builder.write_reconciliation(model, output / "reconciliation.md")
            bid_builder.write_docx(model, output / "proposal.docx")
            self.assertGreater((output / "proposal.docx").stat().st_size, 10000)
            self.assertIn("$4,440.00", (output / "reconciliation.md").read_text())

            from docx import Document

            document = Document(output / "proposal.docx")
            text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            for phrase in (
                "Quote is good for 30 Days",
                "Proposal # Q10001",
                "Per Plans Dated:",
                "NON CREDIT CUSTOMERS:",
                "Toilet Partitions By: Accurate",
                "Qty.",
                "Manufacturer Name",
                "Proposal Expiration & Potential Escalation",
            ):
                self.assertIn(phrase, text)


if __name__ == "__main__":
    unittest.main()
